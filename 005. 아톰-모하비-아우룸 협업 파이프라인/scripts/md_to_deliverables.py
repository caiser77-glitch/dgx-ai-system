#!/usr/bin/env python3
# 005 파이프라인 — AURUM_CONVERT_COMMAND 변환기: draft 마크다운 → HWPX/Docx/PDF 실산출물.
# Created: 2026-07-21 by Antigravity(Claude Opus 4.8)  (2026-07-22 표준 표지 추가)
# 배포기(aurum_deployer)가 DRAFT_PATH, OUTPUT_HWP, OUTPUT_PDF, OUTPUT_DOCX 환경변수로 호출.
# HWPX=hangul_mcp, Docx=python-docx, PDF=fpdf2+Noto CJK KR. 표준 표지(사업명·작성기관·작성연월) 포함.
import os, sys, re, json, logging, datetime

logging.getLogger("fontTools").setLevel(logging.ERROR)  # 폰트 서브셋 경고 억제

DRAFT = os.environ.get("DRAFT_PATH", "")
OUT_HWP = os.environ.get("OUTPUT_HWP", "")
OUT_PDF = os.environ.get("OUTPUT_PDF", "")
OUT_DOCX = os.environ.get("OUTPUT_DOCX", "")
CJK_FONT = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
ORG = "㈜아우룸생태연구소"

def load_body(path):
    t = open(path, encoding="utf-8", errors="ignore").read()
    if t.startswith("---"):               # frontmatter 제거
        end = t.find("\n---", 3)
        if end != -1:
            t = t[end + 4:]
    i = t.find("아우룸 검토 의견")           # 아우룸 내부 검토는 납품물서 제외
    if i != -1:
        t = t[:i].rstrip().rstrip('#').rstrip().rstrip('-').rstrip()
    return t.strip()

def load_cover(draft_path):
    """summary.md 에서 사업명·대상종을 읽어 표준 표지 라인을 만든다."""
    summ = draft_path.replace(".draft.md", ".summary.md")
    proj = species = ""
    if os.path.exists(summ):
        t = open(summ, encoding="utf-8", errors="ignore").read()
        m = re.search(r'project_name:\s*"?([^"\n]+)"?', t); proj = (m.group(1).strip() if m else "")
        m = re.search(r'detected_protected_species:\s*(\[.*\])', t)
        if m:
            try: species = ", ".join(json.loads(m.group(1)))
            except Exception: species = ""
    ym = datetime.date.today().strftime("%Y년 %m월")
    return [
        ("blank", ""), ("blank", ""),
        ("cover_title", proj or "생태조사"),
        ("blank", ""),
        ("cover_sub", "법정보호종 조사 결과 검토 보고서 (초안)"),
        ("blank", ""), ("blank", ""), ("blank", ""),
        ("cover_meta", f"대상 법정보호종 : {species}" if species else "대상 법정보호종 : 본문 참조"),
        ("cover_meta", f"작성기관 : {ORG}"),
        ("cover_meta", f"작 성 : {ym}"),
        ("pagebreak", ""),
    ]

def clean_inline(s):
    s = re.sub(r'[\U0001F000-\U0001FAFF☀-➿←-⇿⌀-⏿]', '', s)  # 이모지/기호 제거
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = re.sub(r'(?<!\*)\*(?!\*)(.+?)\*', r'\1', s)
    s = re.sub(r'`(.+?)`', r'\1', s)
    s = re.sub(r'\[\[?([^\]]+?)\]?\]', r'\1', s)
    return s.strip()

def parse_lines(body):
    out = []
    for raw in body.splitlines():
        s = raw.strip()
        if not s:
            out.append(("blank", "")); continue
        if re.match(r'^-{3,}$', s):
            out.append(("blank", "")); continue
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            out.append((f"h{min(len(m.group(1)), 3)}", clean_inline(m.group(2)))); continue
        m = re.match(r'^[-*+]\s+(.*)', s)
        if m:
            out.append(("bullet", clean_inline(m.group(1)))); continue
        out.append(("para", clean_inline(s)))
    return out

def make_hwpx(lines, path):
    from hangul_mcp.server import create_hwpx_document
    blocks = []
    for k, tx in lines:
        if k == "blank":
            blocks.append({"type": "text", "text": ""}); continue
        if k == "pagebreak":
            blocks.append({"type": "text", "text": "─" * 30}); continue
        prefix = "  • " if k == "bullet" else ""
        if k in ("cover_title", "cover_sub"):
            prefix = ""
        blocks.append({"type": "text", "text": prefix + tx})
    create_hwpx_document(path, json.dumps(blocks, ensure_ascii=False))

def make_docx(lines, path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = docx.Document()
    for k, tx in lines:
        if k == "blank":
            d.add_paragraph(""); continue
        if k == "pagebreak":
            d.add_page_break(); continue
        if k == "cover_title":
            p = d.add_paragraph(); r = p.add_run(tx); r.bold = True; r.font.size = docx.shared.Pt(26)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; continue
        if k == "cover_sub":
            p = d.add_paragraph(); r = p.add_run(tx); r.font.size = docx.shared.Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; continue
        if k == "cover_meta":
            p = d.add_paragraph(tx); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; continue
        if k in ("h1", "h2", "h3"):
            d.add_heading(tx, level=int(k[1]))
        elif k == "bullet":
            d.add_paragraph(tx, style="List Bullet")
        else:
            d.add_paragraph(tx)
    d.save(path)

def make_pdf(lines, path):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    pdf.add_font("noto", "", CJK_FONT)
    for k, tx in lines:
        if k == "blank":
            pdf.ln(4); continue
        if k == "pagebreak":
            pdf.add_page(); continue
        if not tx and k not in ("cover_title", "cover_sub"):
            continue
        try:
            if k == "cover_title":
                pdf.set_font("noto", size=24); pdf.multi_cell(0, 14, tx, align="C")
            elif k == "cover_sub":
                pdf.set_font("noto", size=15); pdf.multi_cell(0, 10, tx, align="C")
            elif k == "cover_meta":
                pdf.set_font("noto", size=12); pdf.multi_cell(0, 8, tx, align="C")
            else:
                size = {"h1": 15, "h2": 13, "h3": 12}.get(k, 10.5)
                pdf.set_font("noto", size=size)
                pdf.multi_cell(pdf.epw, size * 0.62, ("  • " if k == "bullet" else "") + tx)
        except Exception:
            continue
    pdf.output(path)

def main():
    if not (DRAFT and os.path.exists(DRAFT)):
        print(f"DRAFT_PATH 없음: {DRAFT}", file=sys.stderr); sys.exit(2)
    lines = load_cover(DRAFT) + parse_lines(load_body(DRAFT))
    errs = []
    for name, fn, out in (("HWPX", make_hwpx, OUT_HWP), ("DOCX", make_docx, OUT_DOCX), ("PDF", make_pdf, OUT_PDF)):
        try:
            fn(lines, out)
        except Exception as e:
            errs.append(f"{name}:{e}")
    if errs:
        print("일부 변환 실패: " + " | ".join(errs), file=sys.stderr); sys.exit(1)
    print(f"변환완료(표준표지 포함) HWPX/Docx/PDF: {OUT_HWP}")

if __name__ == "__main__":
    main()
